import os
import re
import sys
import helper
import argparse
from typing import Optional
from docx import Document
from docx.oxml.ns import qn


def add_table_footnotes(
    docx_in: str,
    docx_out: str,
    table_dir: str,
    footnotes_yaml: str,
    config_yaml: Optional[str],
    include_object_path: bool = False,
    fail_on_missing_metadata: bool = True,
):
    # Load standard footnotes from a JSON file
    footnotes = helper.load_yaml(footnotes_yaml)

    # load config.yaml or set empty dict for defaults.
    if config_yaml is not None:
        config = helper.load_yaml(config_yaml)
    else:
        config = {}

    document = Document(docx_in)

    # Define magic string pattern that allows for flexible paths
    start_pattern = (
        r"\{rpfy\}\:"  # Matches "{rpfy}:" and any directory structure following it
    )
    end_pattern = r"\.[^.]+$"
    magic_pattern = re.compile(start_pattern + ".*?" + end_pattern)
    paragraphs = document.paragraphs
    missing_metadata = False

    for i, par in enumerate(paragraphs):
        matches = magic_pattern.findall(par.text)
        if not matches:
            continue

        if len(matches) > len(set(matches)):
            print(f"Duplicate table names found in paragraph {i+1}")

        for match in matches:
            # Generalized extraction of the table name
            table_name = match.replace("{rpfy}:", "").strip()
            table_path = os.path.normpath(os.path.join(table_dir, table_name))

            if os.path.exists(table_path):
                metadata = helper.load_metadata(
                    os.path.dirname(table_path),
                    os.path.basename(table_path)
                )

                add_footnote = False
                if metadata is None:
                    missing_metadata = True
                else:
                    add_footnote = True
                    meta_text_dict = helper.create_meta_text_lines(
                        footnotes, metadata, include_object_path, "table", config
                    )

                if add_footnote:
                    # this gets xml index rather than paragraph index
                    current_p = par._p
                    body_elements = list(document.element.body)
                    p_index = body_elements.index(current_p)

                    # w:tbl is directly after matching magic string
                    table = body_elements[p_index + 1]
                    if table.tag == qn("w:tbl"):
                        new_paragraph = helper.create_footnote_paragraph(
                            meta_text_dict, table_name, i, config
                        )
                        document.element.body.insert(
                            document.element.body.index(table) + 1, new_paragraph
                        )

    # Save the processed document
    if missing_metadata and fail_on_missing_metadata:
        print(
            "Output not created due to missing metadata. Please check logs for missing metadata files."
        )
        sys.exit(1)
    else:
        document.save(docx_out)
        print(f"Processed file saved at '{docx_out}'.")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Add table footnotes to input docx document"
    )
    parser.add_argument(
        "-i", "--input", type=str, required=True, help="input docx file path"
    )
    parser.add_argument(
        "-o", "--output", type=str, required=True, help="Path to output docx file"
    )
    parser.add_argument(
        "-d", "--table_dir", type=str, required=True, help="Path to tables directory"
    )
    parser.add_argument(
        "-f",
        "--footnotes",
        type=str,
        required=True,
        help="path to standard footnotes yaml",
    )
    parser.add_argument(
        "-c", "--config", type=str, required=True, help="Path to config.yaml file"
    )
    parser.add_argument(
        "-b",
        "--object",
        type=lambda x: x.lower() in ["true", "t"],
        help="include object path",
    )
    parser.add_argument(
        "-m",
        "--fail_metadata",
        type=lambda x: x.lower() in ["true", "t"],
        help="Allow missing metadata files",
    )
    args = parser.parse_args()

    add_table_footnotes(
        args.input,
        args.output,
        args.table_dir,
        args.footnotes,
        args.config,
        args.object,
        args.fail_metadata,
    )
