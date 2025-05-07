import os
import re

import helper
import tempfile
import argparse
from typing import Optional

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from PIL import Image, ImageDraw, ImageFont

from parse_magic_string import parse_magic_string

def add_figure(
    docx_in: str,
    docx_out: str,
    figure_dir: str,
    config_yaml: Optional[str],
    fig_width: Optional[float] = None,
    fig_height: Optional[float] = None,
):
    document = Document(docx_in)

    # load config.yaml or set empty dict for defaults.
    if config_yaml is not None:
        config = helper.load_yaml(config_yaml)
    else:
        config = {}

    # Define magic string pattern
    # Matches "{rpfy}:" and any directory structure following it
    start_pattern = r"\{rpfy\}\:"  
    end_pattern = r"\.[^.]+$"
    magic_pattern = re.compile(start_pattern + ".*?" + end_pattern)

    found_magic_strings = []
    new_paragraphs = []

    for i, par in enumerate(document.paragraphs):
        matches = magic_pattern.findall(par.text)
        if matches:
            if len(matches) > len(set(matches)):
                print(f"Duplicate figure names found in paragraph {i+1}.")

            for match in matches:
                # Extract the image directory from the match
                # figure_name now contains potentially a list
                # of file names and args. like
                # [file.ext, file2.ext]<width: 4, height: 6>
                figure_args = parse_magic_string(match)

                if len(figure_args) > 1:
                    figures = list(reversed(figure_args.keys()))
                else:
                    figures = list(figure_args.keys())
                
                for fig_idx, figure in enumerate(figures):
                    extension = os.path.splitext(figure)[1].lower()
                    if extension not in [".png", ".csv", ".rds"]: 
                        print(f"Unsupported figure file extension: {figure}. Please save as .png")
                        continue

                    add_label = False
                    if len(figures) > 1 and config.get("label_multi_figures", False):
                        add_label = True

                    found_magic_strings.append(figure)
                    image_path = os.path.join(figure_dir, figure)
                    if os.path.exists(image_path):
                        if add_label:
                            # since list is reversed need to use correct
                            # index, index 0 corresponds to the last
                            # element in list so len(figures) - 1 for 1-index
                            labeled_image = add_label_to_image(
                                image_path, len(figures) - fig_idx - 1
                            )
                        else:
                            labeled_image = image_path

                        # Create new paragraph and add to list
                        new_par = document.add_paragraph()
                        run = new_par.add_run()
                        # +2 because word paragraphs are 1-indexed
                        # and we are adding the paragraph after the
                        # current paragraph
                        new_paragraphs.append((i + 2, new_par))

                        # can only use embedded_size if the args are there
                        # args = {
                        #    'file.ext': {'width': '5', 'height': '8'},
                        #    'file2.ext': {'height': '8'},
                        #    'file3.ext': {'width': '5'},
                        #    'file4.ext': {}
                        # }
                        if config.get("use_embedded_size", True) and set(
                            figure_args[figure].keys()
                        ).intersection(["width", "height"]):
                            run.add_picture(
                                labeled_image,
                                width=(
                                    Inches(float(figure_args[figure].get("width")))
                                    if "width" in figure_args[figure]
                                    else None
                                ),
                                height=(
                                    Inches(float(figure_args[figure].get("height")))
                                    if "height" in figure_args[figure]
                                    else None
                                ),
                            )

                        elif config.get("use_artifact_size", False):
                            run.add_picture(labeled_image)

                        else:
                            default_width = config.get("default_fig_width", 6)
                            if set(figure_args[figure].keys()).intersection(["width", "height"]):
                                run.add_picture(
                                    labeled_image,
                                    width=(
                                        Inches(float(figure_args[figure].get("width")))
                                        if "width" in figure_args[figure]
                                        else None
                                    ),
                                    height=(
                                        Inches(float(figure_args[figure].get("height")))
                                        if "height" in figure_args[figure]
                                        else None
                                    ),
                                )

                            elif fig_width is not None and fig_height is not None:
                                run.add_picture(
                                    labeled_image,
                                    width=Inches(fig_width),
                                    height=Inches(fig_height),
                                )
                            elif fig_width is not None:
                                run.add_picture(labeled_image, width=Inches(fig_width))
                            elif fig_height is not None:
                                run.add_picture(
                                    labeled_image, height=Inches(fig_height)
                                )
                            else:
                                run.add_picture(
                                    labeled_image, width=Inches(default_width)
                                )

                        match config.get("fig_alignment", "center").lower():
                            case "center":
                                new_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            case "left":
                                new_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            case "right":
                                new_par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            case _:
                                new_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Insert new paragraphs
    new_paragraphs.sort(reverse=True, key=lambda x: x[0])
    for par_idx, par in new_paragraphs:
        parent = document.paragraphs[0]._element.getparent()
        par._element.getparent().remove(par._element)
        parent.insert(par_idx, par._element)

    if len(set(found_magic_strings)) != len(found_magic_strings):
        print("Duplicate figure names found in the document.")
    document.save(docx_out)
    print(f"Processed file saved at '{docx_out}'.")


def add_label_to_image(image_path: str, index: int) -> str:
    """
    This function takes in a path to an image and an index
    and adds the corresponding letter to the image upper
    left corner. If the index is > 25 then the label will
    use multiple letters.
    26 = AA
    27 = AB, etc.

    This function saves the updated image to tmp and returns
    the path to the temp image.
    """

    label = helper.create_label(index)

    # load in image and create draw object
    # and set font
    img = Image.open(image_path)
    draw = ImageDraw.Draw(img)
    font = ImageFont.load_default(size=56)

    img_width, img_height = img.size
    # aspect_ratio = img_width / img_height

    left_corner_position = (20, 20)

    draw.rectangle(
        xy=[
            left_corner_position[0] - 5,
            left_corner_position[1] - 5,
            left_corner_position[0] + img_width * 0.025,
            left_corner_position[1] + img_height * 0.025,
        ],
        fill=(255, 255, 255),
    )

    draw.text(left_corner_position, label, fill=(0, 0, 0), font=font)
    temp_file = tempfile.NamedTemporaryFile(
        delete=False, suffix=os.path.splitext(image_path)[1]
    )
    temp_path = temp_file.name
    temp_file.close()

    img.save(temp_path)
    return temp_path


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Add figures to input docx document")
    parser.add_argument(
        "-i", "--input", type=str, required=True, help="input docx file path"
    )
    parser.add_argument("-o", "--output", type=str, required=True, help="output docx")
    parser.add_argument(
        "-d", "--figure_dir", type=str, required=True, help="Path to figures directory"
    )
    parser.add_argument(
        "-c", "--config", type=str, default=None, help="Config yaml path"
    )
    parser.add_argument("-w", "--width", type=str, default=None, help="Figure width")
    parser.add_argument("-g", "--height", type=str, default=None, help="Figure height")
    args = parser.parse_args()

    add_figure(
        args.input, args.output, args.figure_dir, args.config, args.width, args.height
    )
